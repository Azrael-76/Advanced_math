from docx import Document
import json

def parse_docx(input_path):
    document = Document(input_path)
    text = []
    for para in document.paragraphs:
        text.append(para.text)

    # Extracting JSON string from the text
    json_str = None
    for line in text:
        if 'Data' in line:
            json_str = line.strip().split(':')[1]
            break

    # Converting JSON string to dictionary
    if json_str:
        data_dict = json.loads(json_str)
        return data_dict
    else:
        return None

def write_coordinates_to_docx(data_dict, output_path):
    document = Document()
    for item in data_dict['doc_layout']:
        for pos in item['pos']:
            document.add_paragraph(f"x: {pos['x']}, y: {pos['y']}")
        document.add_paragraph('')

    document.save(output_path)

if __name__ == "__main__":
    input_path = r"D:\Coding\Advanced_math\output_raw_docx\page_1.docx"
    output_path = r"D:\Coding\Advanced_math\output_raw_docx\page_1_coor_gpt3.5.docx"

    data_dict = parse_docx(input_path)
    if data_dict:
        write_coordinates_to_docx(data_dict['body'], output_path)
        print("Coordinates written to:", output_path)
    else:
        print("Error: JSON data not found in the document.")
