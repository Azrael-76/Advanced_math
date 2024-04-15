import re
from docx import Document

# 定义正则表达式模式
section_regex = r'(?<=\*\*).+?(?=\*\*)'
question_regex = r'\d+\. (.+?)(?=\d+\. |$)'

# 读取docx文件并提取题目内容
def extract_questions(docx_path):
    document = Document(docx_path)
    text = []
    for paragraph in document.paragraphs:
        text.append(paragraph.text)

    text = "\n".join(text)
    sections = re.findall(section_regex, text, re.DOTALL)

    fill_in_the_blank_questions = []
    essay_questions = []

    for section in sections:
        questions = re.findall(question_regex, section, re.DOTALL)
        if "填空题" in section:
            fill_in_the_blank_questions.extend(questions)
        elif "解答题" in section:
            essay_questions.extend(questions)

    return fill_in_the_blank_questions, essay_questions

# 将提取的题目内容写入新的docx文件
def write_questions(fill_in_the_blank_questions, essay_questions, output_docx_path):
    document = Document()
    document.add_heading('题目如下：', level=1)
    document.add_heading('**一、填空题（本题共5小题，每小题3分，满分15分）**', level=2)
    for i, question in enumerate(fill_in_the_blank_questions, 1):
        document.add_paragraph(f"{i}. {question}")

    document.add_heading('**二、解答题（本题满分8分）**', level=2)
    for i, question in enumerate(essay_questions, 1):
        document.add_paragraph(f"{i}. {question}")

    document.save(output_docx_path)

# 提取题目内容
fill_in_the_blank_questions, essay_questions = extract_questions(r"D:\Coding\Advanced_math\output_raw_docx\page_1.docx")

# 将题目内容写入新的docx文件
write_questions(fill_in_the_blank_questions, essay_questions, r"D:\Coding\Advanced_math\output_raw_docx\page_1_gpt3.5.docx")

print("解析完成并写入文件成功！")
