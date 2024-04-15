from docx import Document

def extract_questions_word(input_docx_path, output_docx_path):
    """
    从Word文档中提取所有从"text":开始到"table_list": []结束的文段。
    如果在遇到"table_list": []之前又遇到了新的"text":，则从新的"text":开始读取，确保每道题只被提取一次。

    参数:
    input_docx_path: 输入的Word文档路径。
    output_docx_path: 输出的Word文档路径。
    """
    doc = Document(input_docx_path)
    extracted_data = []
    current_question = []

    # 查找并提取指定文段
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if '"text":' in text:
            # 如果遇到新的"text":，开始新的提取
            if current_question and '"table_list": []' not in current_question[-1]:
                # 如果在"table_list": []之前又遇到了新的"text":，则从新的"text":开始读取
                current_question = [text]
            else:
                current_question.append(text)
        elif current_question:
            # 添加段落到当前问题
            current_question.append(text)
            if '"table_list": []' in text:
                # 如果遇到"table_list": []，结束当前问题的提取
                extracted_data.append(current_question+['          }'])
                current_question = []

    # 将提取的文段保存到新的Word文档中
    output_doc = Document()
    for question in extracted_data:
        for paragraph_text in question:
            print(paragraph_text)
            output_doc.add_paragraph(paragraph_text)
    output_doc.save(output_docx_path)
    print("*****************************************************************************")
    print("完成截取题目坐标")

# 示例用法
# extract_questions_word(r"D:\Coding\Advanced_math\output_imgs_of_pages\page_1.jpg", r"D:\Coding\Advanced_math\output_cut_docx\page_1.docx")


