import re
from docx import Document

def get_coordinates_part(input_docx_path):
    # 从Word文档中读取文本
    doc = Document(input_docx_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # 匹配"part_title"及其后的第一组四对坐标
    part_pattern = re.compile(r'"part_title".*?(\{\s*"x":\s*(\d+),\s*"y":\s*(\d+)\s*\}\s*,\s*\{\s*"x":\s*(\d+),\s*"y":\s*(\d+)\s*\}\s*,\s*\{\s*"x":\s*(\d+),\s*"y":\s*(\d+)\s*\}\s*,\s*\{\s*"x":\s*(\d+),\s*"y":\s*(\d+)\s*\})', re.DOTALL)

    matches = part_pattern.findall(text)
    all_coordinates = []

    for match in matches:
        coordinates = [{"x": int(match[1]), "y": int(match[2])},
                       {"x": int(match[3]), "y": int(match[4])},
                       {"x": int(match[5]), "y": int(match[6])},
                       {"x": int(match[7]), "y": int(match[8])}]
        all_coordinates.append(coordinates)

    # 打印提取的坐标信息
    print("*****************************************************************************")
    for index, coordinates in enumerate(all_coordinates):
        print(f"Part {index+1}: {coordinates}")
    print("完成提取坐标")  
    return all_coordinates


# 示例用法
# get_coordinates_part(r"D:\Coding\Advanced_math\output_raw_docx\page_1.docx")

import re
from docx import Document

def get_coordinates_sub(input_docx_path):
    # 从Word文档中读取文本
    doc = Document(input_docx_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text

    # 正则表达式匹配"part_title"和"pos_list"
    part_title_pattern = re.compile(r'"part_title".*?(\[.*?\])', re.DOTALL)
    pos_list_pattern = re.compile(r'"pos_list":\[\[({.*?})\]\]', re.DOTALL)

    # 分割文本以"part_title"为界
    parts = part_title_pattern.split(text)[1:]  # 跳过第一个分割结果，因为它可能不包含有效数据

    all_coordinates = []

    for part in parts:
        # 在每个部分中查找"pos_list"
        pos_lists = pos_list_pattern.finditer(part)
        for pos_list_match in pos_lists:
            # 提取坐标数据
            coord_text = pos_list_match.group(1)
            # 匹配四对坐标
            coords = re.findall(r'\{"x":(\d+),"y":(\d+)\}', coord_text)
            if len(coords) == 4:  # 确保匹配到四对坐标
                coordinates = [{"x": int(coord[0]), "y": int(coord[1])} for coord in coords]
                all_coordinates.append(coordinates)

    # 打印提取的坐标信息
    print("*****************************************************************************")
    for index, coordinates in enumerate(all_coordinates):
        print(f"Coordinates {index+1}: {coordinates}")
    
    print("完成提取坐标")  
    return all_coordinates

# 示例用法
# get_coordinates_sub("D:\\Coding\\Advanced_math\\output_cut_docx\\page_1.docx")